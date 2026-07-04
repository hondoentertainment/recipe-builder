"""Export recipes to a formatted Word document."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from recipe_extractor import Recipe


def add_recipe(doc: Document, recipe: Recipe, image_dir: Path | None = None) -> None:
    doc.add_heading(recipe.title, level=1)

    if recipe.description:
        p = doc.add_paragraph(recipe.description)
        p.runs[0].italic = True

    meta_parts = []
    if recipe.servings:
        meta_parts.append(f"Servings: {recipe.servings}")
    if recipe.prep_time:
        meta_parts.append(f"Prep: {recipe.prep_time}")
    if recipe.cook_time:
        meta_parts.append(f"Cook: {recipe.cook_time}")
    if meta_parts:
        meta = doc.add_paragraph(" | ".join(meta_parts))
        meta.runs[0].font.size = Pt(10)
        meta.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    if image_dir and recipe.source_image:
        img_path = image_dir / recipe.source_image
        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    doc.add_heading("Ingredients", level=2)
    for ingredient in recipe.ingredients:
        doc.add_paragraph(ingredient, style="List Bullet")

    doc.add_heading("Instructions", level=2)
    for i, step in enumerate(recipe.instructions, 1):
        doc.add_paragraph(f"{i}. {step}")

    if recipe.notes:
        doc.add_heading("Notes", level=2)
        doc.add_paragraph(recipe.notes)

    doc.add_paragraph()
    doc.add_paragraph("—" * 40)
    doc.add_paragraph()


def export_recipes_to_word(
    recipes: list[Recipe],
    output_path: Path,
    image_dir: Path | None = None,
) -> Path:
    doc = Document()

    title = doc.add_heading("Recipe Collection", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph(f"{len(recipes)} recipe(s) from Google Photos")
    doc.add_paragraph()

    for recipe in recipes:
        add_recipe(doc, recipe, image_dir)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
